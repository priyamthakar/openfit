"""DOCX (Word) report renderer for openfit."""

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:
    import matplotlib.figure

    from openfit.results import FitResult

DISCLAIMER: str = (
    "This report was generated using openfit (open-source). "
    "Results should be independently verified for regulatory or "
    "clinical decision-making."
)


def _fmt(value: float | None) -> str:
    """Format a float to 5 significant figures, or 'N/A'."""
    import math

    if value is None:
        return "N/A"
    try:
        f = float(value)
    except (TypeError, ValueError):
        return "N/A"
    if not math.isfinite(f):
        return str(f)
    return f"{f:.5g}"


def _fig_to_png_bytes(fig: matplotlib.figure.Figure) -> BytesIO:
    """Convert a matplotlib Figure to a BytesIO PNG buffer.

    Parameters
    ----------
    fig : matplotlib.figure.Figure
        Figure to convert.

    Returns
    -------
    BytesIO
        BytesIO buffer containing PNG data, positioned at start.
    """
    import matplotlib.pyplot as plt

    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set solid background shading on a cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    shading.append(shading_elem)


def _make_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a professionally styled table to the document.

    Parameters
    ----------
    doc : Document
        The python-docx Document object.
    headers : list[str]
        Column header texts.
    rows : list[list[str]]
        Data rows (each row is a list of strings).
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Dark header row colors
    HEADER_BG = "1e3a5f"
    ALT_ROW_1 = "ffffff"
    ALT_ROW_2 = "f8fafc"

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_shading(cell, HEADER_BG)

    # Data rows with alternating colors
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        bg = ALT_ROW_1 if row_idx % 2 == 0 else ALT_ROW_2
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)
            if col_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_shading(cell, bg)

    # Add a blank paragraph after the table for spacing
    doc.add_paragraph("")


def render_docx_report(result: FitResult) -> bytes:
    """Generate a DOCX fit report as bytes.

    Parameters
    ----------
    result : FitResult
        A completed fit result.

    Returns
    -------
    bytes
        DOCX file contents as bytes.
    """
    from openfit.plotting import fit_overlay_plot, residual_plot
    from openfit.report.html import _decide_log_x

    doc = Document()

    # --- Document title ---
    title = doc.add_heading("openfit Fit Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Metadata ---
    timestamp = getattr(result.spec, "timestamp", "N/A")
    meta = doc.add_paragraph(f"Model: {result.model_id}  |  {timestamp}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    n_obs = int(result.n_obs)
    n_params = int(result.n_params)
    df = n_obs - n_params

    # ------------------------------------------------------------------
    # Parameter Estimates table
    # ------------------------------------------------------------------
    doc.add_heading("Parameter Estimates", level=1)

    param_headers = ["Parameter", "Value", "SE", "95% CI"]
    param_rows = []
    for name in result.params:
        val = _fmt(result.params[name])
        se_val = _fmt(result.se.get(name) if isinstance(result.se, dict) else None)
        ci_entry = result.ci.get(name) if isinstance(result.ci, dict) else None
        ci_str = f"[{_fmt(ci_entry[0])}, {_fmt(ci_entry[1])}]" if ci_entry is not None else "N/A"
        param_rows.append([name, val, se_val, ci_str])

    _make_table(doc, param_headers, param_rows)

    # ------------------------------------------------------------------
    # Goodness of Fit table
    # ------------------------------------------------------------------
    doc.add_heading("Goodness of Fit", level=1)

    gof_headers = ["Metric", "Value"]
    gof_rows = [
        ["R^2", _fmt(result.r_squared)],
        ["AICc", _fmt(result.aicc)],
        ["BIC", _fmt(result.bic)],
        ["AIC", _fmt(result.aic)],
        ["RSS", _fmt(result.rss)],
        ["Observations (n)", str(n_obs)],
        ["Parameters (p)", str(n_params)],
        ["Degrees of freedom", str(df)],
        ["Weight scheme", str(result.weight_scheme)],
    ]

    _make_table(doc, gof_headers, gof_rows)

    # ------------------------------------------------------------------
    # Optional ROUT outlier table
    # ------------------------------------------------------------------
    rout_result = getattr(result, "rout_result", None)
    if rout_result is not None:
        doc.add_heading("ROUT Outlier Detection", level=1)

        rout_headers = ["Metric", "Value"]
        rout_rows = [
            ["Total points", str(len(result.x))],
            ["Flagged outliers", str(rout_result.n_outliers)],
            ["Q parameter", f"{rout_result.Q * 100:.1f}%"],
        ]

        _make_table(doc, rout_headers, rout_rows)

    # ------------------------------------------------------------------
    # Fit Overlay Plot
    # ------------------------------------------------------------------
    doc.add_heading("Fit Overlay", level=1)

    log_x = _decide_log_x(result.x)

    try:
        overlay_fig = fit_overlay_plot(result, log_x=log_x)
        overlay_buf = _fig_to_png_bytes(overlay_fig)
        doc.add_picture(overlay_buf, width=Inches(5))
        overlay_buf.close()
    except Exception:
        doc.add_paragraph("[Plot unavailable]")

    doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Residual Plot
    # ------------------------------------------------------------------
    doc.add_heading("Residuals", level=1)

    try:
        resid_fig = residual_plot(result)
        resid_buf = _fig_to_png_bytes(resid_fig)
        doc.add_picture(resid_buf, width=Inches(5))
        resid_buf.close()
    except Exception:
        doc.add_paragraph("[Plot unavailable]")

    doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Footer / Disclaimer
    # ------------------------------------------------------------------
    disclaimer = doc.add_paragraph(DISCLAIMER)
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.italic = True

    generated = doc.add_paragraph(f"Generated by openfit  |  {timestamp}")
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.runs[0].font.size = Pt(8)
    generated.runs[0].font.italic = True

    # Serialize document to bytes
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    buf.close()
    return docx_bytes
